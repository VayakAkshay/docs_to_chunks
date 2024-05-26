from docx import Document
import pandas as pd
from general_functions import table_to_chunks, get_splitter, convert_into_markdown


def extract_text_from_docx(doc_path):
    doc = Document(doc_path)
    all_text = ""
    for paragraph in doc.paragraphs:
        all_text += paragraph.text + "\n"
    return all_text


def extract_all_tables_from_docx(doc_path):
    doc = Document(doc_path)

    # Initialize a list to store DataFrames
    all_tables = []

    # Loop through each table in the document
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)

        # Convert the table data into a DataFrame
        df = pd.DataFrame(table_data)

        # Optionally, handle headers (assuming the first row is the header)
        if len(df) > 0:
            df.columns = df.iloc[0]  # Set the first row as header
            df = df[1:]  # Remove the header row from the data
            df.reset_index(drop=True, inplace=True)  # Reset index after removing header

        all_tables.append(df)

    return all_tables


if __name__ == "__main__":
    # Add your desire pdf file
    doc_path = "/Users/akshay/Documents/Projects/Sai_work/all_data-merged.docx"
    token_limit = 500
    text_content = extract_text_from_docx(doc_path)
    print(text_content)
    print("Part 1 Done")
    all_tables = extract_all_tables_from_docx(doc_path)
    chunks = []
    print("Part 2 Done")

    for table in all_tables:
        markdown_table = convert_into_markdown(table)
        table_chunks = table_to_chunks(markdown_table, token_limit)

        chunks = chunks + table_chunks

    print("Part 3 Done")
    text_splitter = get_splitter()

    text_splits = text_splitter([text_content])
    print("Part 4 Done")

    all_chunks = text_splits + chunks
    print("Part 5 Done")

    print(all_chunks)